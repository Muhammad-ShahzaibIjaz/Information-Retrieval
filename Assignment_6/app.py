import os
import string
from docx import Document
from difflib import SequenceMatcher
from collections import defaultdict
from flask import Flask, request, jsonify
from flask_cors import CORS



app = Flask(__name__)
CORS(app)



class SearchEngine:
    def __init__(self, folder_path='Documents'):
        self.folder_path = folder_path
        self.documents, self.search_terms, self.content_membership_degrees, self.title_membership_degrees, self.author_membership_degrees = self.setup_search_engine()

    def setup_search_engine(self):
        extracted_documents = self.document_extractor()
        extracted_content = self.extract_fullContext_from_documents(extracted_documents)
        extracted_titles = self.extract_titles_from_documents(extracted_documents)
        extracted_authors = self.extract_author_from_documents(extracted_documents)
        
        content_index, search_terms = self.build_index(extracted_content)
        title_index, _ = self.build_index(extracted_titles)
        author_index, _ = self.build_index(extracted_authors)

        content_membership_degrees = self.calculate_fuzzy_membership(content_index)
        title_membership_degrees = self.calculate_fuzzy_membership(title_index)
        author_membership_degrees = self.calculate_fuzzy_membership(author_index)

        return extracted_documents, search_terms, content_membership_degrees, title_membership_degrees, author_membership_degrees

    def extract_text_from_documents(self, file_path):
        document = Document(file_path)
        doc_data = {'title': "", 'author': "", "content": ""}

        paragraphs = (p.text.strip() for p in document.paragraphs if p.text.strip())
        for text in paragraphs:
            if not doc_data['title']:
                doc_data['title'] = text
            elif text.startswith("Author:"):
                doc_data['author'] = text[7:].strip()
            else:
                doc_data['content'] += f" {text}"

        return doc_data

    def document_extractor(self):
        documents = []
        for file_name in os.listdir(self.folder_path):
            file_path = os.path.join(self.folder_path, file_name)
            if not (os.path.isfile(file_path) and file_name.endswith('.docx')):
                continue

            doc_data = self.extract_text_from_documents(file_path)
            doc_data['file_name'] = file_name
            documents.append(doc_data)

        return documents

    def extract_fullContext_from_documents(self, documents):
        return [
            f"Title: {doc['title']}\nAuthor: {doc['author']}\nContent: {doc['content']}"
            for doc in documents
        ]

    def extract_titles_from_documents(self, documents):
        return [doc['title'] for doc in documents if doc.get('title')]

    def extract_author_from_documents(self, documents):
        return [doc['author'] for doc in documents if doc.get('author')]

    def preprocess_text(self, text):
        stopwords = {"the", "and", "is", "in", "to", "of", "on", "for", "with", "a", "an", "as", "by", "this", "it", "at", "or", "that"}
        translator = str.maketrans('', '', string.punctuation)
        return [word for word in text.lower().translate(translator).split() if word not in stopwords]

    def build_index(self, documents):
        index = defaultdict(lambda: {'doc_ids': [], 'tf': {}})
        search_terms = defaultdict(int)

        for doc_id, content in enumerate(documents):
            words = self.preprocess_text(content)

            for i in range(len(words)):
                word = words[i]
                if doc_id not in index[word]['tf']:
                    index[word]['tf'][doc_id] = 0
                index[word]['tf'][doc_id] += 1
                search_terms[word] += 1

                for j in range(i + 1, min(i + 5, len(words) + 1)):
                    phrase = " ".join(words[i:j])
                    search_terms[phrase] += 1

                index[word]['doc_ids'].append(doc_id)

        return index, search_terms

    def calculate_similarity(self, term1, term2):
        return SequenceMatcher(None, term1, term2).ratio()

    def find_partial_matches(self, query_term, terms, threshold=0.5):
        matches = {}
        for term in terms:
            similarity = self.calculate_similarity(query_term, term)
            if similarity >= threshold:
                matches[term] = similarity
        return matches

    def calculate_fuzzy_membership(self, index):
        membership = defaultdict(lambda: defaultdict(float))
        for term, data in index.items():
            max_tf = max(data['tf'].values()) if data['tf'] else 1
            for doc_id, tf in data['tf'].items():
                membership[term][doc_id] = tf / max_tf
        return membership

    def process_fuzzy_query(self, query, membership_degrees, fuzziness_threshold=0.30):
        query_terms = self.preprocess_text(query)
        doc_scores = defaultdict(float)
        term_weights = {term: query_terms.count(term) for term in query_terms}

        for term in query_terms:
            if term in membership_degrees:
                for doc_id, membership in membership_degrees[term].items():
                    if membership >= fuzziness_threshold:
                        doc_scores[doc_id] += membership * term_weights[term]
            else:
                partial_matches = self.find_partial_matches(term, membership_degrees.keys(), fuzziness_threshold)
                for match_term, similarity in partial_matches.items():
                    for doc_id, membership in membership_degrees[match_term].items():
                        if membership * similarity >= fuzziness_threshold:
                            doc_scores[doc_id] += membership * term_weights[term] * similarity

        max_score = max(doc_scores.values()) if doc_scores else 1
        for doc_id in doc_scores:
            doc_scores[doc_id] /= max_score

        return doc_scores


    def rank_documents(self, doc_scores, documents):
        ranked_docs = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)
        result = []
        for doc_id, score in ranked_docs:
            doc = documents[doc_id]
            snippet = f"{doc['content'][:180].rsplit(' ', 1)[0]}..." if len(doc['content']) > 100 else doc['content']
            result.append({'title': doc['title'], 'author': doc['author'], 'snippet': snippet, 'score': score, 'file_path': doc['file_name']})
        return result


    def suggest_keywords(self, input_text):
        input_text = input_text.lower()
        suggestions = [term for term in self.search_terms if term.startswith(input_text)]
        return sorted(suggestions, key=lambda x: self.search_terms[x], reverse=True)[:5]


    def search(self, query, membership_degrees, fuzziness_threshold):
        doc_scores = self.process_fuzzy_query(query, membership_degrees, fuzziness_threshold)
        results = self.rank_documents(doc_scores, self.documents)
        return results
    


def extract_docx_content(docx_file):
    doc = Document(docx_file)
    
    content = {
        "title": "",
        "author": "",
        "abstract": "",
        "sections": {}
    }
    current_section = None
    current_subsection = None

    # Flags for extracting title, author, and abstract
    title_extracted = False
    extracted_author = False
    extracted_abstract = False
    abstract_start = False  # Initialize the abstract_start flag

    for para in doc.paragraphs:
        # Extract title: Assume it's the first non-empty paragraph that's not a heading
        if not title_extracted and para.text.strip() and not para.style.name.startswith('Heading'):
            content["title"] = para.text.strip()
            title_extracted = True
            continue  # Skip the title from further processing

        # Extract author: Assume the author's name follows a specific format or keyword
        if not extracted_author and "author" in para.text.lower():
            content["author"] = para.text.strip().replace("Author:", "").replace("author:", "").strip()
            extracted_author = True
            continue

        # Extract abstract (assuming it's labeled as "Abstract" or a similar heading)
        if not extracted_abstract and para.text.strip().lower() == "abstract":
            abstract_start = True  # Start extracting the abstract
            continue
        elif abstract_start:
            content["abstract"] += para.text.strip() + "\n"
            if para.text.strip() == "":
                abstract_start = False  # End of abstract when a blank line is encountered
            continue

        # Check if the paragraph is a heading (usually marked as bold)
        if para.style.name.startswith('Heading'):
            heading_level = int(para.style.name.split()[-1])  # Extract heading level
            heading_text = para.text.strip()

            if heading_level == 1:  # Main section
                current_section = heading_text
                content["sections"][current_section] = ""  # Initialize as an empty string for the section content
                current_subsection = None
            elif heading_level == 2:  # Subsection
                current_subsection = heading_text
                if current_section:  # Ensure current_section is initialized
                    # If the section is not a dictionary, convert it into one
                    if isinstance(content["sections"][current_section], str):  # Main section is currently a string
                        content["sections"][current_section] = {}  # Convert it into a dictionary for subsections
                    content["sections"][current_section][current_subsection] = ""  # Initialize empty string for subsection content
            else:
                continue  # For further levels, you can extend logic here if needed

        else:
            # Otherwise, it's content under the current section or subsection
            if current_section:
                if current_subsection:
                    # Append content to the subsection (it's a string)
                    content["sections"][current_section][current_subsection] += para.text.strip() + "\n"
                else:
                    # Append content to the main section (it's a string)
                    content["sections"][current_section] += para.text.strip() + "\n"

    return content



# Initialize the search engine once when the app starts
search_engine = SearchEngine('Documents')

@app.route('/api/suggestions', methods=['GET'])
def suggestions():
    query = request.args.get('query', '')
    if not query:
        return jsonify({"error": "Query parameter is required"}), 400

    suggestions = search_engine.suggest_keywords(query)
    return jsonify({"suggestions": suggestions})


@app.route('/api/v6/search/author', methods=['GET'])
def search_author():
    query = request.args.get('query', '')
    threshold = float(request.args.get('threshold', 0.6))

    if not query:
        return jsonify({"error": "Query parameter is required"}), 400

    results = search_engine.search(query, search_engine.author_membership_degrees, threshold)
    return jsonify({"results": results})

@app.route('/api/v6/search/title', methods=['GET'])
def search_title():
    query = request.args.get('query', '')
    threshold = float(request.args.get('threshold', 0.6))

    if not query:
        return jsonify({"error": "Query parameter is required"}), 400

    results = search_engine.search(query, search_engine.title_membership_degrees, threshold)
    return jsonify({"results": results})

@app.route('/api/v6/search/content', methods=['GET'])
def search_fulltext():
    query = request.args.get('query', '')
    threshold = float(request.args.get('threshold', 0.6))

    if not query:
        return jsonify({"error": "Query parameter is required"}), 400

    results = search_engine.search(query, search_engine.content_membership_degrees, threshold)
    return jsonify({"results": results})

@app.route('/api/v6/article/extract', methods=['GET'])
def extract_content():
    try:
        query = request.args.get('query', '')
        if not query:
            return jsonify({"error": "Query parameter is required"}), 400
        
        data = extract_docx_content('Documents/' + query)
        return jsonify(data), 200
    except Exception as e:
        return jsonify({"error" : str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True)